from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

#218d5e (Green) 33,141,94

# Create a new Document
doc = Document()

# Set margins (in inches)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Function to add a new line
def add_newline(doc, count=1):
    for _ in range(count):
        doc.add_paragraph()

# Function to set font style for paragraphs
def set_font(run, font_name, font_size, color_rgb=None, bold=False, underline=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    if underline:
        run.font.underline = underline

#Add spaces before the start of the title to justify text vertically
add_newline(doc, count=7)

#Title
year = doc.add_paragraph()
year_run = year.add_run('2018')
year_run.italic = True
set_font(year_run, 'Century Schoolbook', 16, color_rgb= (128, 128, 128))
year.alignment = WD_ALIGN_PARAGRAPH.CENTER

title = doc.add_paragraph()
title_run = title.add_run('3rd Quarter Report')
set_font(title_run, 'Century Schoolbook', 39, color_rgb= (33, 141, 94), bold=True)  # Blue color
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

#Add spaces after the end of the title to justify text vertically
add_newline(doc, count=11)

#Project Details
#Project
p1 = doc.add_paragraph()
p1_run_1 = p1.add_run("Project: ")
set_font(p1_run_1, 'Avenir Next', 14, color_rgb= (33, 141, 94), bold=True)
p1_run_2 = p1.add_run('Dualisation of Ibadan-Illorin Road Section II')
set_font(p1_run_2, 'Avenir Next', 14, color_rgb= (0, 0, 0), bold=False)

#Ministry
p2 = doc.add_paragraph()
p2_run_1 = p2.add_run('Ministry: ')
set_font(p2_run_1, 'Avenir Next', 14, color_rgb= (33, 141, 94), bold=True)
p2_run_2 = p2.add_run('Federal Ministry of Works and Housing')
set_font(p2_run_2, 'Avenir Next', 14, color_rgb= (0, 0, 0), bold=False)

#Respondent
p3 = doc.add_paragraph()
p3_run_1 = p3.add_run('Respondent: ')
set_font(p3_run_1, 'Avenir Next', 14, color_rgb= (33, 141, 94), bold=True)
p3_run_2 = p3.add_run('Dr. Akintola Ashimewu')
set_font(p3_run_2, 'Avenir Next', 14, color_rgb= (0, 0, 0), bold=False)

#Date
p4 = doc.add_paragraph()
p4_run_1 = p4.add_run('Date: ')
set_font(p4_run_1, 'Avenir Next', 14, color_rgb= (33, 141, 94), bold=True)
p4_run_2 = p4.add_run('24/09/2024')
set_font(p4_run_2, 'Avenir Next', 14, color_rgb= (0, 0, 0), bold=False)


# Section: General Information
general_info_heading = doc.add_paragraph()
general_info_heading_run = general_info_heading.add_run('GENERAL INFORMATION')
set_font(general_info_heading_run, 'Century Schoolbook', 18, color_rgb=(33, 141, 94), bold=True, underline=WD_UNDERLINE.SINGLE)
general_info_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

general_info_text = ('The project’s major output is 52.00 Km of dual carriageway which is a section '
                     'of Ibadan – Illorin road, connecting the northern part of the country to the '
                     'Lagos port and the South-West.\n\n'
                     'The project commenced in June 2010 with a planned completion date of '
                     'October 2013, the contract is still ongoing.\n\n'
                     'The project was contracted to Messrs. Reynolds Construction Company '
                     'Nigeria Limited. The project traverses towns between Oyo and Ogbomoso, '
                     'Oyo State along the coordinates N 07◦57’58.8, E 004◦03’52.8.')

general_info_para = doc.add_paragraph(general_info_text)
set_font(general_info_para.runs[0], 'Avenir Next', 13)

# Adding spacing
add_newline(doc, 2)


# Add a new section with a single column layout
new_section = doc.add_section(WD_SECTION.CONTINUOUS)

# Change the new section to a single column layout
columns = new_section._sectPr.xpath('./w:cols')[0]
columns.set(qn('w:num'), '1')

###################################################################################################################################
# Add a section with two columns
section = doc.sections[-1]
columns = section._sectPr.xpath('./w:cols')[0]
columns.set(qn('w:num'), '2')

#Bullet List for priority areas and project objectives
priority_area_title = doc.add_paragraph()
priority_area_title_run = priority_area_title.add_run('Targeted Priority Areas')
set_font(priority_area_title_run, 'Avenir Next', 13, color_rgb= (0, 0, 0), bold=True, underline=WD_UNDERLINE.SINGLE)
priority = doc.add_paragraph('Item 1', style='List Bullet')
priority = doc.add_paragraph('Item 2', style='List Bullet')
priority = doc.add_paragraph('Item 3', style='List Bullet')


project_obj_title = doc.add_paragraph()
project_obj_title_run = project_obj_title.add_run('Project Objectives')
set_font(project_obj_title_run, 'Avenir Next', 13, color_rgb= (0, 0, 0), bold=True, underline=WD_UNDERLINE.SINGLE)
objective = doc.add_paragraph('Item 1', style='List Bullet')
objective = doc.add_paragraph('Item 1', style='List Bullet')
objective = doc.add_paragraph('Item 1', style='List Bullet')

# Add a new section with a single column layout
new_section = doc.add_section(WD_SECTION.CONTINUOUS)

# Change the new section to a single column layout
columns = new_section._sectPr.xpath('./w:cols')[0]
columns.set(qn('w:num'), '1')
####################################################################################################################################

# Section: Situational Analysis
situational_analysis_heading = doc.add_paragraph()
situational_analysis_heading_run = situational_analysis_heading.add_run('SITUATIONAL ANALYSIS')
set_font(situational_analysis_heading_run, 'Century Schoolbook', 18, color_rgb=(33, 141, 94), bold=True, underline=WD_UNDERLINE.SINGLE)  # Blue underline
situational_analysis_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Financial Performance
financial_performance_heading = doc.add_paragraph()
financial_performance_heading_run = financial_performance_heading.add_run('Financial Performance')
set_font(financial_performance_heading_run, 'Avenir Next', 13, bold=True, underline=WD_UNDERLINE.SINGLE)  # Underline

financial_performance_text = ('The initial estimated cost for the project was forty-seven billion five hundred '
                              'and four million one hundred and thirty-eight thousand three hundred and '
                              'forty-four-Naira twenty kobo (₦47,504,138,344.20). No amount has been '
                              'released for the project from the 2018 capital expenditure.\n\n'
                              'Twenty-six billion seven hundred and eighty-six million eight hundred and forty '
                              'thousand seven hundred and sixty-three Naira thirty-eight kobo (₦26,786, '
                              '840,763.38) has been committed to the project from inception.')

financial_performance_para = doc.add_paragraph(financial_performance_text)
set_font(financial_performance_para.runs[0], 'Avenir Next', 13)

# Results Delivery - Field Observation
results_heading = doc.add_paragraph()
results_heading_run = results_heading.add_run('Results Delivery - Field Observation')
set_font(results_heading_run, 'Avenir Next', 13, bold=True, underline=WD_UNDERLINE.SINGLE)  # Underline

results_text = ('The project is ongoing at 59% completion with 19.64 km both northern and southern '
                'bound at asphaltic binder course level.\n\n'
                'Most of the work achieved was due to funding from the Sukuk bond. The project is '
                'currently being reviewed to increase the work scope; this is aimed at improving the '
                'load-bearing ability of the road.')

results_para = doc.add_paragraph(results_text)
set_font(results_para.runs[0], 'Avenir Next', 13)


# Section: Challenges
challenges_heading = doc.add_paragraph()
challenges_heading_run = challenges_heading.add_run('Challenges')
set_font(challenges_heading_run, 'Avenir Next', 13, bold=True, underline=WD_UNDERLINE.SINGLE)  # Underline

challenges_text = ('Poor funding from the Federal Government has hampered the progress of the '
                   'project, commuters and trucks have to ply the old Oyo-Ogbomoso Road which is '
                   'always prone to traffic logjams lasting up to six days at a stretch. Also, the hard '
                   'shoulders of the old road are eroded for long stretches at various sections of the road '
                   'constituting a danger to commuters that ply this road.')

challenges_para = doc.add_paragraph(challenges_text)
set_font(challenges_para.runs[0], 'Avenir Next', 13)

### SECTION: Two column Layout for images.

# Add a new section with a single column layout
new_section = doc.add_section(WD_SECTION.CONTINUOUS)

# Change the new section to a single column layout
columns = new_section._sectPr.xpath('./w:cols')[0]
columns.set(qn('w:num'), '1')

# Add a section with two columns
section = doc.sections[-1]
columns = section._sectPr.xpath('./w:cols')[0]
columns.set(qn('w:num'), '2')


#Pictures
doc.add_picture('pictures/1691515191331.jpg', width=Inches(2), height=Inches(3))
doc.add_picture('pictures/1691516252800.jpg', width=Inches(2), height=Inches(3))

# Add a new section with a single column layout
new_section = doc.add_section(WD_SECTION.CONTINUOUS)

# Change the new section to a single column layout
columns = new_section._sectPr.xpath('./w:cols')[0]
columns.set(qn('w:num'), '1')

# Adding spacing
add_newline(doc, 3)

# Section: Recommendations
recommendations_heading = doc.add_paragraph()
recommendations_heading_run = recommendations_heading.add_run('Recommendations')
set_font(recommendations_heading_run, 'Avenir Next', 13, color_rgb=(0, 0, 0), bold=True, underline=WD_UNDERLINE.SINGLE)  # Blue underline

recommendations_text = ('Poor funding from the Federal Government has hampered the progress of the '
                        'project, commuters and trucks have to ply the old Oyo-Ogbomoso Road which is '
                        'always prone to traffic logjams lasting up to six days at a stretch. Also, the hard '
                        'shoulders of the old road are eroded for long stretches at various sections of the road '
                        'constituting a danger to commuters that ply this road.')

recommendations_para = doc.add_paragraph(recommendations_text)
set_font(recommendations_para.runs[0], 'Avenir Next', 13)

# Adding footer
section = doc.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
footer_run = footer_para.add_run('Projects Progress Report - 2018 Q3')
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer_run, 'Avenir Next', 10, color_rgb=(100, 100, 100))

doc.save('generated_report.docx')



