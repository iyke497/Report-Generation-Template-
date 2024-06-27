from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new document
doc = Document()

# Add a section with two columns
section = doc.sections[-1]
columns = section._sectPr.xpath('./w:cols')[0]
columns.set(qn('w:num'), '2')

# Add content to the two-column section
doc.add_paragraph("This is the first column of the two-column layout. " * 5)
doc.add_paragraph("This is the second column of the two-column layout. " * 5)

# Add a new section with a single column layout
new_section = doc.add_section(WD_SECTION.CONTINUOUS)

# Change the new section to a single column layout
columns = new_section._sectPr.xpath('./w:cols')[0]
columns.set(qn('w:num'), '1')

# Add content to the single-column section
doc.add_paragraph("This is the single-column layout for the remainder of the page. " * 5)

# Save the document
doc.save('two_column_single_column_layout.docx')
print("Document saved as 'two_column_single_column_layout.docx'")
